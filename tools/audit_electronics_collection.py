#!/usr/bin/env python3
"""Audit the Replacor electronics article collection before web integration."""

from __future__ import annotations

import argparse
import json
import re
import unicodedata
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn


SAFETY_TERMS = (
    "peligro", "advertencia", "precaución", "precaucion", "seguridad",
    "alta tensión", "alta tension", "bus dc", "bus de continua",
    "condensador cargado", "descarga", "osciloscopio", "masa flotante",
    "transformador de aislamiento", "sonda diferencial",
)
MEASUREMENT_TERMS = (
    "multímetro", "multimetro", "osciloscopio", "pinza", "lcr", "medir",
    "medición", "medicion", "tensión", "tension", "corriente", "resistencia",
    "frecuencia", "forma de onda",
)
REPAIR_TERMS = (
    "avería", "averia", "diagnóstico", "diagnostico", "comprobación",
    "comprobacion", "sustitución", "sustitucion", "falsa avería",
    "falsa averia", "causa", "síntoma", "sintoma",
)


def normalize(value: str) -> str:
    value = "".join(
        char for char in unicodedata.normalize("NFD", value)
        if unicodedata.category(char) != "Mn"
    ).lower()
    return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9]+", " ", value)).strip()


def heading_level(paragraph: Any) -> int | None:
    style = (paragraph.style.name if paragraph.style else "") or ""
    match = re.search(r"(?:Heading|Título|Titulo)\s*([1-6])", style, re.I)
    if match:
        return int(match.group(1))
    p_pr = paragraph._p.pPr
    outline = p_pr.find(qn("w:outlineLvl")) if p_pr is not None else None
    return int(outline.get(qn("w:val"))) + 1 if outline is not None else None


def count_terms(text: str, terms: tuple[str, ...]) -> dict[str, int]:
    haystack = normalize(text)
    result: dict[str, int] = {}
    for term in terms:
        count = haystack.count(normalize(term))
        if count:
            result[term] = count
    return result


def audit_docx(path: Path) -> dict[str, Any]:
    document = Document(path)
    paragraphs = [re.sub(r"\s+", " ", p.text).strip() for p in document.paragraphs]
    paragraphs = [text for text in paragraphs if text]
    headings = []
    for paragraph in document.paragraphs:
        text = re.sub(r"\s+", " ", paragraph.text).strip()
        level = heading_level(paragraph)
        if text and level:
            headings.append({"level": level, "text": text})
    tables = []
    for table in document.tables:
        rows = [[re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells] for row in table.rows]
        if rows:
            tables.append({"rows": len(rows), "columns": max(len(row) for row in rows)})
    with zipfile.ZipFile(path) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/") and not name.endswith("/")]
        links_xml = ""
        try:
            links_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8", errors="ignore")
        except KeyError:
            pass
        external_links = re.findall(r'Target="(https?://[^\"]+)"', links_xml)
    full_text = "\n".join(paragraphs)
    plain_urls = re.findall(r"https?://[^\s<>()]+", full_text)
    possible_sources = [
        text for text in paragraphs
        if re.search(r"\b(?:fuentes?|bibliograf[ií]a|referencias?|datasheet|application note|manual)\b", text, re.I)
    ]
    placeholders = [
        text for text in paragraphs
        if re.search(r"\b(?:pendiente|por completar|por revisar|TODO|XXX|insertar|a confirmar)\b", text, re.I)
    ]
    repeated = Counter(normalize(text) for text in paragraphs if len(text) > 50)
    duplicates = [key for key, count in repeated.items() if count > 1]
    return {
        "file": path.name,
        "paragraphs": len(paragraphs),
        "words": len(full_text.split()),
        "headings": headings,
        "tables": len(tables),
        "table_shapes": tables,
        "embedded_media": len(media),
        "external_links": external_links,
        "plain_urls": plain_urls,
        "source_mentions": possible_sources[-12:],
        "safety_terms": count_terms(full_text, SAFETY_TERMS),
        "measurement_terms": count_terms(full_text, MEASUREMENT_TERMS),
        "repair_terms": count_terms(full_text, REPAIR_TERMS),
        "placeholders": placeholders,
        "repeated_paragraphs": duplicates,
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    reports = []
    for path in sorted(args.source.rglob("*.docx")):
        report = audit_docx(path)
        report["folder"] = path.parent.name
        reports.append(report)
    payload = {
        "articles": len(reports),
        "totals": {
            "words": sum(item["words"] for item in reports),
            "tables": sum(item["tables"] for item in reports),
            "media": sum(item["embedded_media"] for item in reports),
            "headings": sum(len(item["headings"]) for item in reports),
        },
        "reports": reports,
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(payload["totals"], ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
